#!/usr/bin/env python3
"""
Build a reflowable Kindle eBook (EPUB + DOCX) and a Kindle-sized cover
from the paperback manuscript of "Trey: A New Beginning".

Source data:
  - composer-project.json          (text overlays for each page)
  - public/images/books/trey-book-01/
      cover.png                    (paperback cover illustration, square)
      front-01.png .. front-04.png (title / dedication / copyright / quote art)
      page-01.png .. page-13.png   (story illustrations)
      (pages 14-21 are caregiver-guide text only; no illustration needed)

Output (written to ./kindle/):
  - book.epub   reflowable EPUB3 (KDP-compatible)
  - book.docx   Word document (backup / editing)
  - cover.jpg   Kindle front cover (1600x2560, 1.6:1)

Design notes:
  * Print-only artifacts (fixed pages, overlays, text-on-image, page numbers,
    absolute positioning) are stripped.
  * Each story page becomes:     <image/> + <paragraph/>   (block flow)
  * Caregiver guide pages become pure reflowable text (<h2/<h3/<p/<ul).
  * Images are down-sampled to max 1600px longest side, JPEG q=82, to keep
    the EPUB lean for Kindle devices.
"""

from __future__ import annotations

import io
import json
import os
import sys
from pathlib import Path
from typing import Iterable

from PIL import Image
from ebooklib import epub
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
ROOT = Path(__file__).resolve().parent.parent
COMPOSER_JSON = ROOT / "composer-project.json"
ASSETS_DIR = ROOT / "public" / "images" / "books" / "trey-book-01"
OUT_DIR = ROOT / "kindle"
OUT_DIR.mkdir(parents=True, exist_ok=True)
WORK_DIR = OUT_DIR / ".build"
WORK_DIR.mkdir(parents=True, exist_ok=True)

# ---------------------------------------------------------------------------
# Book metadata
# ---------------------------------------------------------------------------
BOOK_TITLE = "Trey: A New Beginning"
BOOK_SUBTITLE = "Book One of the Trey Series"
AUTHOR = "Stephen Marrero"
PUBLISHER = "Lifestyle Creations"
LANGUAGE = "en"
ISBN = "urn:uuid:trey-book-01-kindle"  # placeholder EPUB identifier
COPYRIGHT_YEAR = "2026"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------
def clean_text(raw: str) -> str:
    """Collapse the print-era soft line breaks into clean reflowable text."""
    if not raw:
        return ""
    # Composer stored hard line-breaks for layout. We want flowing prose.
    return " ".join(line.strip() for line in raw.splitlines() if line.strip())


def prepare_image(
    src: Path,
    dest: Path,
    max_side: int = 1600,
    quality: int = 82,
    background=(255, 255, 255),
) -> Path:
    """Resize + JPEG-encode an illustration so the EPUB stays slim."""
    with Image.open(src) as im:
        if im.mode in ("RGBA", "LA", "P"):
            bg = Image.new("RGB", im.size, background)
            rgba = im.convert("RGBA")
            bg.paste(rgba, mask=rgba.split()[-1] if rgba.mode == "RGBA" else None)
            im = bg
        else:
            im = im.convert("RGB")
        w, h = im.size
        scale = min(max_side / max(w, h), 1.0)
        if scale < 1.0:
            im = im.resize((int(w * scale), int(h * scale)), Image.LANCZOS)
        im.save(dest, "JPEG", quality=quality, optimize=True, progressive=True)
    return dest


def build_kindle_cover(src_png: Path, dest_jpg: Path) -> Path:
    """
    Build a Kindle-recommended 1600 x 2560 (1.6:1) cover from the square
    paperback illustration by letterboxing onto an ocean-blue canvas that
    matches the artwork edges.  No spine, no back panel.
    """
    target_w, target_h = 1600, 2560
    with Image.open(src_png) as im:
        im = im.convert("RGB")

        # sample edge color for a seamless extension
        top_row = im.crop((0, 0, im.width, 1)).resize((1, 1)).getpixel((0, 0))
        bot_row = im.crop((0, im.height - 1, im.width, im.height)).resize((1, 1)).getpixel((0, 0))

        # scale illustration to fit inside target, keep aspect, leave room top/bottom
        scale = min(target_w / im.width, target_h / im.height)
        art_w = int(im.width * scale)
        art_h = int(im.height * scale)
        art = im.resize((art_w, art_h), Image.LANCZOS)

        # average of top/bottom sample colors for background
        bg = tuple(int((top_row[c] + bot_row[c]) / 2) for c in range(3))
        canvas = Image.new("RGB", (target_w, target_h), bg)

        # paint a subtle vertical gradient from top-edge to bottom-edge color
        gradient = Image.new("RGB", (1, target_h))
        for y in range(target_h):
            t = y / (target_h - 1)
            gradient.putpixel(
                (0, y),
                tuple(int(top_row[c] * (1 - t) + bot_row[c] * t) for c in range(3)),
            )
        canvas.paste(gradient.resize((target_w, target_h)), (0, 0))

        # center the illustration
        offset = ((target_w - art_w) // 2, (target_h - art_h) // 2)
        canvas.paste(art, offset)

        canvas.save(dest_jpg, "JPEG", quality=90, optimize=True, progressive=True)
    return dest_jpg


# ---------------------------------------------------------------------------
# Load manuscript
# ---------------------------------------------------------------------------
with COMPOSER_JSON.open() as f:
    composer = json.load(f)

pages = composer["books"][0]["pages"]

# Build a clean {pageNumber: text} map from composer overlays
page_text: dict[int, str] = {}
for page in pages:
    pn = page.get("pageNumber")
    if pn is None:
        continue
    overlays = page.get("textOverlays") or []
    joined = "\n".join(clean_text(o.get("text", "")) for o in overlays).strip()
    page_text[pn] = joined

# ---------------------------------------------------------------------------
# Section definitions (reflowable structure)
# ---------------------------------------------------------------------------
# Story pages: illustration + prose paragraph (no overlay text).
story_pages = []
for pn in range(1, 14):
    story_pages.append(
        {
            "page_number": pn,
            "image": ASSETS_DIR / f"page-{pn:02d}.png",
            "text": page_text.get(pn, ""),
        }
    )

# Caregiver guide (pages 14-21) — pure text, structured with headings + lists.
# We split each page's raw text into a title + body based on the source layout.
caregiver_sections = [
    {
        "title": "For Parents and Caregivers",
        "body_html": (
            "<p>The following pages offer gentle guidance to support "
            "conversations about loss, adoption, and belonging as you read "
            "Trey&rsquo;s story together.</p>"
        ),
    },
    {
        "title": "A Note to Caregivers",
        "body_html": (
            "<p>Children may not always have the words to express grief, "
            "loss, or change &mdash; but they feel it deeply.</p>"
            "<p>As you read Trey&rsquo;s story together, your role is not to "
            "fix the feelings, but to stay present with them.</p>"
            "<p>Your calm presence helps your child feel safe.</p>"
            "<p>There are no perfect words &mdash; connection is what "
            "matters most.</p>"
        ),
    },
    {
        "title": "How to Use This Book",
        "body_html": (
            "<ul>"
            "<li>Read at your child&rsquo;s pace</li>"
            "<li>Pause when emotions arise</li>"
            "<li>Allow repeated questions</li>"
            "<li>It&rsquo;s okay if your child does not want to talk right away</li>"
            "</ul>"
            "<p>Understanding happens gradually and often returns in layers.</p>"
        ),
    },
    {
        "title": "Pages 1&ndash;2: Loss &amp; Separation",
        "body_html": (
            "<p>Trey loses his parents and is suddenly alone.</p>"
            "<h3>Your child may feel:</h3>"
            "<ul><li>Sadness</li><li>Fear</li><li>Questions about safety</li></ul>"
            "<h3>What to say:</h3>"
            "<ul>"
            "<li>&ldquo;Trey feels really sad.&rdquo;</li>"
            "<li>&ldquo;That is a really big change.&rdquo;</li>"
            "<li>&ldquo;It makes sense his heart feels heavy.&rdquo;</li>"
            "</ul>"
        ),
    },
    {
        "title": "Pages 3&ndash;5: Safe Adults &amp; Support",
        "body_html": (
            "<p>Raya and Marina step in to help Trey.</p>"
            "<h3>Your child learns:</h3>"
            "<ul>"
            "<li>Adults can care and protect</li>"
            "<li>Help exists after loss</li>"
            "</ul>"
            "<h3>What to say:</h3>"
            "<ul>"
            "<li>&ldquo;Raya wants Trey to feel safe.&rdquo;</li>"
            "<li>&ldquo;There are people who help children.&rdquo;</li>"
            "</ul>"
        ),
    },
    {
        "title": "Pages 6&ndash;10: A New Family &amp; Belonging",
        "body_html": (
            "<p>Trey is invited into a new home.</p>"
            "<h3>Your child may feel:</h3>"
            "<ul><li>Excited</li><li>Nervous</li><li>Unsure</li></ul>"
            "<h3>What to say:</h3>"
            "<ul>"
            "<li>&ldquo;New things can feel exciting and scary.&rdquo;</li>"
            "<li>&ldquo;It&rsquo;s okay to take time to feel safe.&rdquo;</li>"
            "<li>&ldquo;Families grow together.&rdquo;</li>"
            "</ul>"
        ),
    },
    {
        "title": "Pages 7 &amp; 11&ndash;13: Differences, Healing &amp; Mixed Feelings",
        "body_html": (
            "<p>Cray experiences the world in his own way, and Trey learns "
            "that love can grow in unexpected places.</p>"
            "<h3>What your child learns:</h3>"
            "<ul>"
            "<li>Differences are safe and valuable</li>"
            "<li>You can feel more than one thing at once</li>"
            "<li>New love does not replace old love</li>"
            "</ul>"
            "<h3>What to say:</h3>"
            "<ul>"
            "<li>&ldquo;Trey feels happy and sad.&rdquo;</li>"
            "<li>&ldquo;Both feelings are okay.&rdquo;</li>"
            "</ul>"
        ),
    },
    {
        "title": "Supporting Your Child",
        "body_html": (
            "<h3>Helpful responses:</h3>"
            "<ul>"
            "<li>&ldquo;I&rsquo;m here with you.&rdquo;</li>"
            "<li>&ldquo;That makes sense.&rdquo;</li>"
            "<li>&ldquo;You&rsquo;re safe.&rdquo;</li>"
            "<li>&ldquo;You can always talk to me.&rdquo;</li>"
            "</ul>"
            "<h3>Try to avoid:</h3>"
            "<ul>"
            "<li>&ldquo;Don&rsquo;t be sad&rdquo;</li>"
            "<li>&ldquo;You&rsquo;re fine&rdquo;</li>"
            "<li>&ldquo;Be strong&rdquo;</li>"
            "</ul>"
            "<h3>Closing Message</h3>"
            "<p>Love can grow again, even after loss.</p>"
            "<p>With your presence, your child can learn they are safe, "
            "cared for, and not alone.</p>"
        ),
    },
]


# ---------------------------------------------------------------------------
# Prepare images (compress for Kindle)
# ---------------------------------------------------------------------------
print("Preparing images...")
prepared_story: dict[int, Path] = {}
for sp in story_pages:
    out = WORK_DIR / f"page-{sp['page_number']:02d}.jpg"
    prepare_image(sp["image"], out)
    prepared_story[sp["page_number"]] = out
    print(f"  page-{sp['page_number']:02d}.jpg")

# Kindle cover
print("Building Kindle cover (1600x2560)...")
cover_src = ASSETS_DIR / "cover.png"
cover_jpg = OUT_DIR / "cover.jpg"
build_kindle_cover(cover_src, cover_jpg)

# small cover copy for embedding inside the EPUB (Kindle uses the external
# cover for the bookshelf thumbnail, and the internal cover page for in-book)
epub_cover = WORK_DIR / "cover.jpg"
with Image.open(cover_jpg) as im:
    im.save(epub_cover, "JPEG", quality=85, optimize=True, progressive=True)


# ---------------------------------------------------------------------------
# Build EPUB
# ---------------------------------------------------------------------------
print("Building EPUB...")
book = epub.EpubBook()
book.set_identifier(ISBN)
book.set_title(BOOK_TITLE)
book.set_language(LANGUAGE)
book.add_author(AUTHOR)
book.add_metadata("DC", "publisher", PUBLISHER)
book.add_metadata("DC", "rights", f"Copyright \u00a9 {COPYRIGHT_YEAR} {AUTHOR}. All rights reserved.")
book.add_metadata("DC", "description", (
    "When a tiny octopus named Trey loses his parents, the ocean feels "
    "impossibly big and quiet. But through the kindness of caring helpers "
    "and the warmth of a new family, Trey discovers that love doesn\u2019t "
    "replace what was lost \u2014 it grows around it."
))

# Cover
with open(epub_cover, "rb") as f:
    book.set_cover("cover.jpg", f.read())

# Stylesheet (reflowable, Kindle-friendly)
css = """
@namespace epub "http://www.idpf.org/2007/ops";
html, body { margin: 0; padding: 0; }
body {
  font-family: Georgia, "Times New Roman", serif;
  line-height: 1.5;
  text-align: left;
  margin: 0 0.5em;
}
h1, h2, h3 { font-family: "Helvetica Neue", Helvetica, Arial, sans-serif; }
h1 { font-size: 1.6em; margin: 1.2em 0 0.6em; text-align: center; }
h2 { font-size: 1.3em; margin: 1em 0 0.5em; text-align: center; }
h3 { font-size: 1.05em; margin: 0.8em 0 0.3em; }
p { margin: 0 0 0.8em; text-indent: 0; }
p.story {
  font-size: 1.05em;
  line-height: 1.55;
  text-align: left;
  margin: 0.6em 0 1.4em;
}
ul { margin: 0.4em 0 0.8em 1.1em; padding: 0; }
li { margin: 0.2em 0; }
.center { text-align: center; }
figure.illustration {
  margin: 0 0 1em;
  padding: 0;
  text-align: center;
  page-break-inside: avoid;
}
figure.illustration img {
  display: block;
  max-width: 100%;
  height: auto;
  margin: 0 auto;
}
.title-page { text-align: center; margin-top: 2em; }
.title-page .title { font-size: 2em; font-weight: bold; margin: 0.3em 0; }
.title-page .subtitle { font-size: 1.1em; color: #555; margin: 0.2em 0 1.5em; }
.title-page .author { font-size: 1.2em; margin: 2em 0 0; }
.title-page .publisher { font-size: 0.9em; color: #666; margin-top: 3em; }
.copyright { font-size: 0.9em; color: #333; margin-top: 3em; }
.copyright p { margin: 0.4em 0; }
.dedication { font-style: italic; text-align: center; margin-top: 4em; }
"""
css_item = epub.EpubItem(
    uid="stylesheet",
    file_name="styles/stylesheet.css",
    media_type="text/css",
    content=css.encode("utf-8"),
)
book.add_item(css_item)


def make_chapter(file_name: str, title: str, body_html: str) -> epub.EpubHtml:
    ch = epub.EpubHtml(
        title=title,
        file_name=file_name,
        lang=LANGUAGE,
    )
    ch.content = (
        f"<html xmlns=\"http://www.w3.org/1999/xhtml\" xmlns:epub=\"http://www.idpf.org/2007/ops\" lang=\"{LANGUAGE}\">"
        f"<head><title>{title}</title>"
        f"<link rel=\"stylesheet\" type=\"text/css\" href=\"styles/stylesheet.css\"/></head>"
        f"<body>{body_html}</body></html>"
    )
    ch.add_item(css_item)
    return ch


# Add image items to the EPUB
for pn, path in prepared_story.items():
    with open(path, "rb") as f:
        data = f.read()
    img = epub.EpubItem(
        uid=f"img_page_{pn:02d}",
        file_name=f"images/page-{pn:02d}.jpg",
        media_type="image/jpeg",
        content=data,
    )
    book.add_item(img)


# Title page
title_html = (
    "<section class='title-page'>"
    f"<h1 class='title'>{BOOK_TITLE}</h1>"
    f"<p class='subtitle'>{BOOK_SUBTITLE}</p>"
    f"<p class='author'>{AUTHOR}</p>"
    f"<p class='publisher'>{PUBLISHER}</p>"
    "</section>"
)
title_ch = make_chapter("title.xhtml", BOOK_TITLE, title_html)
book.add_item(title_ch)

# Copyright
copyright_html = (
    "<section class='copyright'>"
    f"<h2>Copyright</h2>"
    f"<p>{BOOK_TITLE}</p>"
    f"<p>Copyright &copy; {COPYRIGHT_YEAR} {AUTHOR}. All rights reserved.</p>"
    "<p>No part of this publication may be reproduced, distributed, or "
    "transmitted in any form or by any means, including photocopying, "
    "recording, or other electronic or mechanical methods, without the prior "
    "written permission of the publisher, except in the case of brief "
    "quotations embodied in critical reviews and certain other "
    "noncommercial uses permitted by copyright law.</p>"
    f"<p>Published by {PUBLISHER}.</p>"
    "<p>This is a work of fiction. Names, characters, places, and incidents "
    "are the product of the author&rsquo;s imagination.</p>"
    "<p>First Kindle eBook edition.</p>"
    "</section>"
)
copyright_ch = make_chapter("copyright.xhtml", "Copyright", copyright_html)
book.add_item(copyright_ch)

# Dedication
dedication_html = (
    "<section class='dedication'>"
    "<h2>Dedication</h2>"
    "<p>For every child learning that love can grow again, "
    "and for the families who hold space for them.</p>"
    "</section>"
)
dedication_ch = make_chapter("dedication.xhtml", "Dedication", dedication_html)
book.add_item(dedication_ch)

# Story
story_chapters = []
for sp in story_pages:
    pn = sp["page_number"]
    text = sp["text"].replace("\n", " ")
    body = (
        "<section class='story'>"
        f"<figure class='illustration'>"
        f"<img src='images/page-{pn:02d}.jpg' alt='Illustration for page {pn}'/>"
        f"</figure>"
        f"<p class='story'>{text}</p>"
        "</section>"
    )
    ch = make_chapter(f"story-{pn:02d}.xhtml", f"Chapter {pn}", body)
    book.add_item(ch)
    story_chapters.append(ch)

# Caregiver guide
caregiver_chapters = []
for idx, sec in enumerate(caregiver_sections, start=1):
    body = f"<section class='caregiver'><h2>{sec['title']}</h2>{sec['body_html']}</section>"
    ch = make_chapter(f"caregiver-{idx:02d}.xhtml", sec["title"], body)
    book.add_item(ch)
    caregiver_chapters.append(ch)


# Navigation
book.toc = (
    (epub.Section("Front Matter"), (title_ch, copyright_ch, dedication_ch)),
    (epub.Section("Story"), tuple(story_chapters)),
    (epub.Section("For Parents &amp; Caregivers"), tuple(caregiver_chapters)),
)
book.add_item(epub.EpubNcx())
book.add_item(epub.EpubNav())

# Spine (reading order)
book.spine = [
    "nav",
    title_ch,
    copyright_ch,
    dedication_ch,
    *story_chapters,
    *caregiver_chapters,
]

# Reflowable flag (makes clear to Kindle Previewer)
book.add_metadata(None, "meta", "", {"name": "original-resolution", "content": "none"})
book.add_metadata(None, "meta", "", {"name": "book-type", "content": "children"})
book.add_metadata(None, "meta", "", {"name": "primary-writing-mode", "content": "horizontal-tb"})
book.add_metadata(None, "meta", "", {"name": "fixed-layout", "content": "false"})

epub_path = OUT_DIR / "book.epub"
epub.write_epub(str(epub_path), book, {})
print(f"  wrote {epub_path}")


# ---------------------------------------------------------------------------
# Build DOCX
# ---------------------------------------------------------------------------
print("Building DOCX...")
doc = Document()

# Default body style
style = doc.styles["Normal"]
style.font.name = "Georgia"
style.font.size = Pt(12)

def docx_heading(text: str, level: int = 1, center: bool = True):
    h = doc.add_heading(text, level=level)
    if center:
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER

def docx_para(text: str, center: bool = False, italic: bool = False, size: int | None = None):
    p = doc.add_paragraph()
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    if italic:
        run.italic = True
    if size is not None:
        run.font.size = Pt(size)
    return p

# Title page
docx_heading(BOOK_TITLE, level=0, center=True)
docx_para(BOOK_SUBTITLE, center=True, italic=True)
doc.add_paragraph()
docx_para(AUTHOR, center=True, size=14)
doc.add_paragraph()
docx_para(PUBLISHER, center=True, size=11)
doc.add_page_break()

# Copyright
docx_heading("Copyright", level=1, center=True)
docx_para(BOOK_TITLE, center=True)
docx_para(f"Copyright © {COPYRIGHT_YEAR} {AUTHOR}. All rights reserved.", center=True)
docx_para(
    "No part of this publication may be reproduced, distributed, or transmitted "
    "in any form or by any means, including photocopying, recording, or other "
    "electronic or mechanical methods, without the prior written permission of "
    "the publisher, except in the case of brief quotations embodied in critical "
    "reviews and certain other noncommercial uses permitted by copyright law."
)
docx_para(f"Published by {PUBLISHER}.", center=True)
docx_para(
    "This is a work of fiction. Names, characters, places, and incidents are "
    "the product of the author’s imagination.",
    center=True,
)
docx_para("First Kindle eBook edition.", center=True)
doc.add_page_break()

# Dedication
docx_heading("Dedication", level=1, center=True)
docx_para(
    "For every child learning that love can grow again, and for the families "
    "who hold space for them.",
    center=True,
    italic=True,
)
doc.add_page_break()

# Story
for sp in story_pages:
    pn = sp["page_number"]
    img_path = prepared_story[pn]
    img_para = doc.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = img_para.add_run()
    run.add_picture(str(img_path), width=Inches(5.5))
    body = doc.add_paragraph(sp["text"].replace("\n", " "))
    body.alignment = WD_ALIGN_PARAGRAPH.LEFT
    body.paragraph_format.space_after = Pt(12)

doc.add_page_break()

# Caregiver sections — replicate content cleanly as text
caregiver_text_sections = [
    ("For Parents and Caregivers",
     ["The following pages offer gentle guidance to support conversations "
      "about loss, adoption, and belonging as you read Trey’s story together."]),
    ("A Note to Caregivers", [
        "Children may not always have the words to express grief, loss, or change — but they feel it deeply.",
        "As you read Trey’s story together, your role is not to fix the feelings, but to stay present with them.",
        "Your calm presence helps your child feel safe.",
        "There are no perfect words — connection is what matters most.",
    ]),
    ("How to Use This Book", [
        ("bullet", "Read at your child’s pace"),
        ("bullet", "Pause when emotions arise"),
        ("bullet", "Allow repeated questions"),
        ("bullet", "It’s okay if your child does not want to talk right away"),
        "Understanding happens gradually and often returns in layers.",
    ]),
    ("Pages 1–2: Loss & Separation", [
        "Trey loses his parents and is suddenly alone.",
        ("h3", "Your child may feel:"),
        ("bullet", "Sadness"),
        ("bullet", "Fear"),
        ("bullet", "Questions about safety"),
        ("h3", "What to say:"),
        ("bullet", "“Trey feels really sad.”"),
        ("bullet", "“That is a really big change.”"),
        ("bullet", "“It makes sense his heart feels heavy.”"),
    ]),
    ("Pages 3–5: Safe Adults & Support", [
        "Raya and Marina step in to help Trey.",
        ("h3", "Your child learns:"),
        ("bullet", "Adults can care and protect"),
        ("bullet", "Help exists after loss"),
        ("h3", "What to say:"),
        ("bullet", "“Raya wants Trey to feel safe.”"),
        ("bullet", "“There are people who help children.”"),
    ]),
    ("Pages 6–10: A New Family & Belonging", [
        "Trey is invited into a new home.",
        ("h3", "Your child may feel:"),
        ("bullet", "Excited"),
        ("bullet", "Nervous"),
        ("bullet", "Unsure"),
        ("h3", "What to say:"),
        ("bullet", "“New things can feel exciting and scary.”"),
        ("bullet", "“It’s okay to take time to feel safe.”"),
        ("bullet", "“Families grow together.”"),
    ]),
    ("Pages 7 & 11–13: Differences, Healing & Mixed Feelings", [
        "Cray experiences the world in his own way, and Trey learns that love can grow in unexpected places.",
        ("h3", "What your child learns:"),
        ("bullet", "Differences are safe and valuable"),
        ("bullet", "You can feel more than one thing at once"),
        ("bullet", "New love does not replace old love"),
        ("h3", "What to say:"),
        ("bullet", "“Trey feels happy and sad.”"),
        ("bullet", "“Both feelings are okay.”"),
    ]),
    ("Supporting Your Child", [
        ("h3", "Helpful responses:"),
        ("bullet", "“I’m here with you.”"),
        ("bullet", "“That makes sense.”"),
        ("bullet", "“You’re safe.”"),
        ("bullet", "“You can always talk to me.”"),
        ("h3", "Try to avoid:"),
        ("bullet", "“Don’t be sad”"),
        ("bullet", "“You’re fine”"),
        ("bullet", "“Be strong”"),
        ("h3", "Closing Message"),
        "Love can grow again, even after loss.",
        "With your presence, your child can learn they are safe, cared for, and not alone.",
    ]),
]

for title, blocks in caregiver_text_sections:
    docx_heading(title, level=1, center=True)
    for block in blocks:
        if isinstance(block, tuple):
            kind, txt = block
            if kind == "bullet":
                doc.add_paragraph(txt, style="List Bullet")
            elif kind == "h3":
                docx_heading(txt, level=2, center=False)
        else:
            doc.add_paragraph(block)
    doc.add_paragraph()

docx_path = OUT_DIR / "book.docx"
doc.save(str(docx_path))
print(f"  wrote {docx_path}")

# Clean up intermediate work directory
import shutil
shutil.rmtree(WORK_DIR, ignore_errors=True)

print("\nDone.")
print(f"  {OUT_DIR/'book.epub'}")
print(f"  {OUT_DIR/'book.docx'}")
print(f"  {OUT_DIR/'cover.jpg'}")
